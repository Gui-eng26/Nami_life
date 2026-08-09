"""
nami_identidade.py — Fonte única de verdade da identidade visual da Nami em documentos.

Princípio de arquitetura aplicado: SINGLE WRITE-POINT.
Nenhum relatório da Nami define cor, fonte ou espaçamento por conta própria.
Todo documento passa por este módulo. Se a identidade mudar, muda aqui — e só aqui.

Uso:
    python nami_identidade.py entrada.docx saida.docx \
        --titulo "MODELO DE NEGÓCIO" \
        --subtitulo "Adesão a tratamentos médicos por agente de IA no WhatsApp" \
        --contexto "Entregável da Trilha 3 — Venture" \
        --org "Mescla Empreende — PUC Campinas" \
        --autor "Guilherme — fundador" \
        --data "Entrega: 07/08/2026  ·  Apresentação presencial: 10/08/2026" \
        --rodape "Nami Life · Modelo de Negócio"

Dependências: python-docx, Pillow (opcional), zipfile/re (stdlib).
"""

import argparse
import os
import re
import shutil
import zipfile
import tempfile

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# =============================================================================
# 1. TOKENS DE MARCA  (única definição autorizada)
# =============================================================================

TOKENS = {
    # --- Núcleo: extraído do branding kit oficial da Nami Life ---
    "laranja":          "FC4C02",  # primária oficial (RGB 252,76,2)
    "laranja_escuro":   "C43C00",  # texto pequeno sobre branco (contraste AA)
    "laranja_tint":     "FFEDE4",  # fundo de destaque
    "off_white":        "F6FFFF",  # secundária oficial do kit (RGB 246,255,255)

    # --- Apoio aprovado ---
    "marinho":          "0F2B46",  # azul marinho — títulos e cabeçalho de tabela
    "marinho_tint":     "E9F0F6",
    "verde_whatsapp":   "128C7E",  # verde WhatsApp escuro (legível como texto)
    "verde_tint":       "E7F5F1",

    # --- Neutros ---
    "texto":            "1A2430",  # grafite — corpo de texto
    "texto_secundario": "5B6B7A",
    "linha":            "D9E1E8",
    "fundo_suave":      "F3F6F8",  # zebra de tabela
    "branco":           "FFFFFF",

    # --- Estados epistêmicos (o que distingue um documento da Nami) ---
    "dado_verificado":  "0F2B46",  # marinho
    "dado_tint":        "E9F0F6",
    "hipotese":         "A8710A",  # âmbar
    "hipotese_tint":    "FDF3E0",
    "alerta":           "A32A1E",  # tijolo (harmoniza com o laranja)
    "alerta_tint":      "FBECEC",
    "decisao":          "0F7A5A",  # verde
    "decisao_tint":     "E6F4EF",
}

FONTE_DOCUMENTO = "Arial"   # Tier "Entrega" — ver GUIDANCE, seção Tipografia


# =============================================================================
# 2. MAPA DE MIGRAÇÃO — paleta antiga (verde) -> paleta Nami
#    Mantém a SEMÂNTICA já existente nos documentos e troca apenas o pigmento.
# =============================================================================

MAPA_CORES = {
    "1e4d3e": TOKENS["marinho"],           # H1 / cabeçalho de tabela
    "2e6b55": TOKENS["laranja_escuro"],    # H2
    "1a2420": TOKENS["texto"],             # corpo
    "5a6b64": TOKENS["texto_secundario"],  # texto de apoio
    "1f6fb5": TOKENS["dado_verificado"],   # nota informativa
    "9b2c2c": TOKENS["alerta"],            # advertência / lacuna
    "8a6a15": TOKENS["hipotese"],          # hipótese
    "1e6b45": TOKENS["decisao"],           # decisão / validado
}

MAPA_FUNDOS = {
    "1e4d3e": TOKENS["marinho"],           # fill do cabeçalho de tabela
    "f4f6f5": TOKENS["fundo_suave"],       # zebra
    "e8f0ec": TOKENS["decisao_tint"],      # callout de decisão
    "fbf3e0": TOKENS["hipotese_tint"],     # callout de hipótese
    "e6eff7": TOKENS["dado_tint"],         # callout de nota
}


# =============================================================================
# 3. Camada 1 — reescrita determinística de tokens no XML
# =============================================================================

def _remapear_tokens_xml(caminho_docx: str, destino_docx: str) -> dict:
    """Troca cores e fundos no document.xml preservando 100% do conteúdo.

    Retorna o relatório de substituições (auditoria)."""
    relatorio = {}
    with tempfile.TemporaryDirectory() as tmp:
        with zipfile.ZipFile(caminho_docx) as z:
            z.extractall(tmp)

        alvo = os.path.join(tmp, "word", "document.xml")
        xml = open(alvo, encoding="utf-8").read()

        for antiga, nova in MAPA_CORES.items():
            padrao = re.compile(r'(<w:color w:val=")' + antiga + r'(")', re.I)
            xml, n = padrao.subn(r"\g<1>" + nova + r"\g<2>", xml)
            if n:
                relatorio[f"color {antiga}->{nova}"] = n

        for antiga, nova in MAPA_FUNDOS.items():
            padrao = re.compile(r'(w:fill=")' + antiga + r'(")', re.I)
            xml, n = padrao.subn(r"\g<1>" + nova + r"\g<2>", xml)
            if n:
                relatorio[f"fill {antiga}->{nova}"] = n

        open(alvo, "w", encoding="utf-8").write(xml)

        if os.path.exists(destino_docx):
            os.remove(destino_docx)
        with zipfile.ZipFile(destino_docx, "w", zipfile.ZIP_DEFLATED) as z:
            for raiz, _, arquivos in os.walk(tmp):
                for a in arquivos:
                    completo = os.path.join(raiz, a)
                    z.write(completo, os.path.relpath(completo, tmp))
    return relatorio


# =============================================================================
# 4. Helpers de baixo nível
# =============================================================================

def _borda(paragrafo, lado="bottom", cor=None, tamanho=12, espaco=6):
    """Aplica borda a um parágrafo (usada como filete de marca)."""
    cor = cor or TOKENS["laranja"]
    pPr = paragrafo._p.get_or_add_pPr()
    pBdr = pPr.find(qn("w:pBdr"))
    if pBdr is None:
        pBdr = OxmlElement("w:pBdr")
        pPr.append(pBdr)
    el = OxmlElement(f"w:{lado}")
    el.set(qn("w:val"), "single")
    el.set(qn("w:sz"), str(tamanho))
    el.set(qn("w:space"), str(espaco))
    el.set(qn("w:color"), cor)
    pBdr.append(el)


def _quebra_de_pagina_antes(paragrafo, ativar=True):
    pPr = paragrafo._p.get_or_add_pPr()
    el = pPr.find(qn("w:pageBreakBefore"))
    if el is None:
        el = OxmlElement("w:pageBreakBefore")
        pPr.append(el)
    el.set(qn("w:val"), "1" if ativar else "0")


def _campo(paragrafo, instrucao, cor, tamanho_pt, negrito=False):
    """Insere um campo Word (ex.: PAGE, NUMPAGES)."""
    run = paragrafo.add_run()
    ini = OxmlElement("w:fldChar"); ini.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve")
    instr.text = instrucao
    fim = OxmlElement("w:fldChar"); fim.set(qn("w:fldCharType"), "end")
    run._r.append(ini); run._r.append(instr); run._r.append(fim)
    run.font.size = Pt(tamanho_pt)
    run.font.name = FONTE_DOCUMENTO
    run.font.color.rgb = RGBColor.from_string(cor)
    run.bold = negrito
    return run


def _formatar(run, texto=None, tamanho=10, cor=None, negrito=False,
              italico=False, caixa_alta=False):
    if texto is not None:
        run.text = texto
    run.font.name = FONTE_DOCUMENTO
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONTE_DOCUMENTO)
    run.font.size = Pt(tamanho)
    run.font.color.rgb = RGBColor.from_string(cor or TOKENS["texto"])
    run.bold = negrito
    run.italic = italico
    if caixa_alta:
        el = OxmlElement("w:caps"); el.set(qn("w:val"), "1")
        run._element.rPr.append(el)
    return run


def _paragrafo_limpo(p, alinhamento=WD_ALIGN_PARAGRAPH.CENTER,
                     antes=0, depois=0, entrelinha=None):
    p.alignment = alinhamento
    pf = p.paragraph_format
    pf.space_before = Pt(antes)
    pf.space_after = Pt(depois)
    if entrelinha:
        pf.line_spacing = entrelinha
    return p


def _inserir_antes_de_elemento(elemento, dono):
    """Insere um novo paragrafo imediatamente antes de um elemento do corpo
    (paragrafo OU tabela). Necessario para posicionar blocos entre tabelas."""
    novo = OxmlElement("w:p")
    elemento.addprevious(novo)
    from docx.text.paragraph import Paragraph
    return Paragraph(novo, dono)


def _inserir_paragrafo_antes(paragrafo_ref):
    novo = OxmlElement("w:p")
    paragrafo_ref._p.addprevious(novo)
    from docx.text.paragraph import Paragraph
    return Paragraph(novo, paragrafo_ref._parent)


# =============================================================================
# 5. Camada 2 — capa, sumário, cabeçalho e rodapé
# =============================================================================

def aplicar_capa(doc, logo, titulo, subtitulo, contexto, org, autor, data):
    """Monta a capa padrao da Nami sobre os paragrafos de capa ja existentes.

    Slots esperados (layout de capa Nami):
      p0 marca | p1 titulo | p2 subtitulo | [filete] | p3 contexto
      p4 organizacao | p5 autor | p6 data | p7 vazio (quebra de pagina)
    """
    p = list(doc.paragraphs)

    # p0: assinatura visual (marca circular)
    _paragrafo_limpo(p[0], antes=80, depois=16)
    p[0].add_run().add_picture(logo, height=Cm(3.4))

    def _reescrever(par, texto, tamanho, cor, negrito=False, caixa=False,
                    antes=0, depois=4):
        for r in list(par.runs):
            r._element.getparent().remove(r._element)
        _paragrafo_limpo(par, antes=antes, depois=depois)
        if texto:
            _formatar(par.add_run(), texto, tamanho=tamanho, cor=cor,
                      negrito=negrito, caixa_alta=caixa)

    _reescrever(p[1], titulo, 26, TOKENS["marinho"], negrito=True, depois=4)
    _reescrever(p[2], subtitulo, 12, TOKENS["texto_secundario"], depois=16)

    filete = _inserir_paragrafo_antes(p[3])
    _paragrafo_limpo(filete, antes=0, depois=16)
    _borda(filete, "bottom", TOKENS["laranja"], tamanho=18, espaco=1)
    _formatar(filete.add_run(), "", tamanho=2)

    _reescrever(p[3], contexto, 13, TOKENS["laranja_escuro"],
                negrito=True, caixa=True, depois=8)
    _reescrever(p[4], org, 11, TOKENS["marinho"], negrito=True, depois=4)
    _reescrever(p[5], autor, 10, TOKENS["texto_secundario"], depois=3)
    _reescrever(p[6], data, 10, TOKENS["texto_secundario"], depois=0)
    _reescrever(p[7], None, 10, TOKENS["texto_secundario"], depois=0)

    # quebra de pagina ao final da capa
    p[7].add_run().add_break(WD_BREAK.PAGE)
    return p[7]


def aplicar_sumario(doc, titulos, fim_da_capa=None):
    """Sumario estatico posicionado logo apos a capa.

    Sem numeros de pagina: nao exige atualizar campo (F9) e nunca desatualiza."""
    fim_da_capa = fim_da_capa if fim_da_capa is not None else doc.paragraphs[7]
    ancora = fim_da_capa._p.getnext()    # primeiro elemento apos a capa
    if ancora is None:
        return
    dono = fim_da_capa._parent

    cab = _inserir_antes_de_elemento(ancora, dono)
    _paragrafo_limpo(cab, WD_ALIGN_PARAGRAPH.LEFT, antes=0, depois=10)
    _formatar(cab.add_run(), "SUM\u00c1RIO", tamanho=16,
              cor=TOKENS["marinho"], negrito=True)
    _borda(cab, "bottom", TOKENS["laranja"], tamanho=12, espaco=6)

    for t in titulos:
        linha = _inserir_antes_de_elemento(ancora, dono)
        _paragrafo_limpo(linha, WD_ALIGN_PARAGRAPH.LEFT, antes=0, depois=6)
        _formatar(linha.add_run(), t, tamanho=11, cor=TOKENS["texto"])

    nota = _inserir_antes_de_elemento(ancora, dono)
    _paragrafo_limpo(nota, WD_ALIGN_PARAGRAPH.LEFT, antes=18, depois=0)
    _borda(nota, "left", TOKENS["laranja"], tamanho=18, espaco=8)
    _formatar(nota.add_run(),
              "  Conven\u00e7\u00e3o de leitura \u2014 marinho: dado verificado em fonte "
              "prim\u00e1ria \u00b7 \u00e2mbar: hip\u00f3tese a validar \u00b7 "
              "tijolo: advert\u00eancia ou lacuna registrada \u00b7 "
              "verde: decis\u00e3o tomada \u00b7 laranja: proposta de valor.",
              tamanho=8.5, cor=TOKENS["texto_secundario"], italico=True)

    quebra = _inserir_antes_de_elemento(ancora, dono)
    _paragrafo_limpo(quebra, WD_ALIGN_PARAGRAPH.LEFT, antes=0, depois=0)
    quebra.add_run().add_break(WD_BREAK.PAGE)


def aplicar_cabecalho_rodape(doc, logo, rodape_texto):
    sec = doc.sections[0]
    sec.different_first_page_header_footer = True  # capa limpa

    # ---- Cabeçalho ----
    h = sec.header
    h.is_linked_to_previous = False
    ph = h.paragraphs[0]
    for r in list(ph.runs):
        r._element.getparent().remove(r._element)
    ph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    ph.paragraph_format.space_after = Pt(4)
    ph.add_run().add_picture(logo, height=Cm(0.52))
    _formatar(ph.add_run(), "  NAMI LIFE", tamanho=9,
              cor=TOKENS["marinho"], negrito=True)
    _formatar(ph.add_run(), f"   ·   {rodape_texto}", tamanho=9,
              cor=TOKENS["texto_secundario"])
    _borda(ph, "bottom", TOKENS["linha"], tamanho=6, espaco=4)

    # ---- Rodapé ----
    f = sec.footer
    f.is_linked_to_previous = False
    pf = f.paragraphs[0]
    for r in list(pf.runs):
        r._element.getparent().remove(r._element)
    pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _formatar(pf.add_run(), "Nami Life   ·   ", tamanho=8.5,
              cor=TOKENS["texto_secundario"])
    _campo(pf, " PAGE ", TOKENS["laranja_escuro"], 8.5, negrito=True)
    _formatar(pf.add_run(), " de ", tamanho=8.5, cor=TOKENS["texto_secundario"])
    _campo(pf, " NUMPAGES ", TOKENS["texto_secundario"], 8.5)

    # Rodapé da capa: apenas o filete, sem numeração
    ff = sec.first_page_footer
    ff.is_linked_to_previous = False
    pff = ff.paragraphs[0]
    for r in list(pff.runs):
        r._element.getparent().remove(r._element)
    pff.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _formatar(pff.add_run(),
              "A solução é digital, mas o cuidado nunca foi tão humano.",
              tamanho=9, cor=TOKENS["laranja_escuro"], italico=True)


def aplicar_hierarquia(doc):
    """Padroniza H1/H2/H3 e reforça a leitura por seção."""
    primeiro_h1 = True
    for p in doc.paragraphs:
        estilo = (p.style.name or "").lower()
        if estilo == "heading 1":
            _borda(p, "bottom", TOKENS["laranja"], tamanho=12, espaco=6)
            p.paragraph_format.space_before = Pt(0 if primeiro_h1 else 18)
            p.paragraph_format.space_after = Pt(10)
            if not primeiro_h1:
                _quebra_de_pagina_antes(p, True)
            primeiro_h1 = False
            for r in p.runs:
                if r.text.strip():
                    r.font.size = Pt(17)
        elif estilo == "heading 2":
            p.paragraph_format.space_before = Pt(14)
            p.paragraph_format.space_after = Pt(5)
        elif estilo == "heading 3":
            p.paragraph_format.space_before = Pt(11)
            p.paragraph_format.space_after = Pt(4)


def normalizar_paginacao(doc):
    """Remove quebras de pagina manuais que ficaram redundantes.

    Principio: a paginacao passa a ser propriedade da HIERARQUIA
    (pageBreakBefore no Titulo 1), nao de quebras posicionadas a mao.
    So sao removidas as quebras imediatamente anteriores a um Titulo 1 —
    as demais sao preservadas por serem decisoes editoriais do autor.
    """
    paragrafos = list(doc.paragraphs)
    removidas = 0
    for i, p in enumerate(paragrafos[:-1]):
        if p.text.strip():
            continue
        brs = [b for b in p._p.findall(".//" + qn("w:br"))
               if b.get(qn("w:type")) == "page"]
        if not brs:
            continue
        seguinte = paragrafos[i + 1]
        if (seguinte.style.name or "").lower() == "heading 1":
            for b in brs:
                b.getparent().remove(b)
            removidas += 1
    return removidas


def aplicar_callouts(doc):
    """Callouts de 1 celula: barra lateral + fundo na cor do seu estado epistemico.

    O estado NAO e adivinhado: e derivado do fundo ja existente no documento
    (que carrega a semantica original) e apenas refinado por palavra-chave
    quando o documento distingue advertencia de hipotese dentro do mesmo tom.
    """
    POR_FUNDO = {
        "E9F0F6": ("dado_verificado", "dado_tint"),
        "FDF3E0": ("hipotese", "hipotese_tint"),
        "E6F4EF": ("decisao", "decisao_tint"),
        "F3F6F8": ("texto_secundario", "fundo_suave"),
    }
    ALERTA = ("advert", "lacuna")
    registro = []

    for i, t in enumerate(doc.tables):
        if len(t.rows) != 1 or len(t.columns) != 1:
            continue
        celula = t.rows[0].cells[0]
        tcPr = celula._tc.get_or_add_tcPr()
        shd = tcPr.find(qn("w:shd"))
        fundo = (shd.get(qn("w:fill")) if shd is not None else "") or ""
        titulo = (celula.paragraphs[0].text or "").strip().lower()

        chave_cor, chave_tint = POR_FUNDO.get(fundo.upper(),
                                              ("dado_verificado", "dado_tint"))

        # refinamentos deterministicos
        if any(k in titulo[:40] for k in ALERTA):
            chave_cor, chave_tint = "alerta", "alerta_tint"
        elif titulo.startswith("decis"):
            chave_cor, chave_tint = "decisao", "decisao_tint"
        elif "a nami ajuda pessoas" in titulo:
            chave_cor, chave_tint = "laranja", "laranja_tint"

        cor, tint = TOKENS[chave_cor], TOKENS[chave_tint]

        if shd is None:
            shd = OxmlElement("w:shd")
            tcPr.append(shd)
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), tint)

        antigo = tcPr.find(qn("w:tcBorders"))
        if antigo is not None:
            tcPr.remove(antigo)
        bordas = OxmlElement("w:tcBorders")
        for lado in ("top", "left", "bottom", "right"):
            el = OxmlElement(f"w:{lado}")
            if lado == "left":
                el.set(qn("w:val"), "single")
                el.set(qn("w:sz"), "30")
                el.set(qn("w:color"), cor)
            else:
                el.set(qn("w:val"), "nil")
            el.set(qn("w:space"), "0")
            bordas.append(el)
        tcPr.append(bordas)
        registro.append((i, chave_cor, titulo[:45]))
    return registro


# =============================================================================
# 6. Orquestração
# =============================================================================

def aplicar_identidade(entrada, saida, logo, titulo, subtitulo, contexto,
                       org, autor, data, rodape, sumario=None):
    tmp = saida + ".tmp.docx"
    relatorio = _remapear_tokens_xml(entrada, tmp)

    doc = Document(tmp)
    fim_capa = aplicar_capa(doc, logo, titulo, subtitulo, contexto,
                            org, autor, data)
    if sumario:
        aplicar_sumario(doc, sumario, fim_da_capa=fim_capa)
    aplicar_cabecalho_rodape(doc, logo, rodape)
    aplicar_hierarquia(doc)
    relatorio["quebras manuais removidas"] = normalizar_paginacao(doc)
    aplicar_callouts(doc)
    doc.save(saida)
    os.remove(tmp)
    return relatorio


def _cli():
    ap = argparse.ArgumentParser(description="Aplica a identidade visual da Nami a um .docx")
    ap.add_argument("entrada")
    ap.add_argument("saida")
    ap.add_argument("--logo", default="nami_marca_circular_900.png")
    ap.add_argument("--titulo", required=True)
    ap.add_argument("--subtitulo", default="")
    ap.add_argument("--contexto", default="")
    ap.add_argument("--org", default="")
    ap.add_argument("--autor", default="")
    ap.add_argument("--data", default="")
    ap.add_argument("--rodape", default="Nami Life")
    args = ap.parse_args()
    rel = aplicar_identidade(args.entrada, args.saida, args.logo, args.titulo,
                             args.subtitulo, args.contexto, args.org,
                             args.autor, args.data, args.rodape)
    for k, v in rel.items():
        print(f"  {k}: {v}")
    print("OK ->", args.saida)


if __name__ == "__main__":
    _cli()
